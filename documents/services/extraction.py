"""
Document Text Extraction Service.
Extracts text from various document formats.
"""

import os
from typing import Optional, Tuple
from django.core.files.uploadedfile import UploadedFile


class TextExtractionService:
    """
    Belge metin çıkarma servisi.
    PDF, Word ve diğer formatlardan metin çıkarır.
    """
    
    def extract_text(self, file_path: str) -> Tuple[str, dict]:
        """
        Dosyadan metin çıkar.
        
        Args:
            file_path: Dosya yolu
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        if not os.path.exists(file_path):
            return "", {"error": "Dosya bulunamadı"}
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return self._extract_from_pdf(file_path)
        elif file_ext in ['.doc', '.docx']:
            return self._extract_from_word(file_path)
        elif file_ext == '.txt':
            return self._extract_from_txt(file_path)
        elif file_ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
            return self._extract_from_image(file_path)
        else:
            return "", {"error": f"Desteklenmeyen format: {file_ext}"}
    
    def extract_from_uploaded_file(self, uploaded_file: UploadedFile) -> Tuple[str, dict]:
        """
        UploadedFile nesnesinden metin çıkar.
        
        Args:
            uploaded_file: Django UploadedFile
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        file_ext = os.path.splitext(uploaded_file.name)[1].lower()
        content = uploaded_file.read()
        uploaded_file.seek(0)  # Reset file pointer
        
        if file_ext == '.pdf':
            return self._extract_from_pdf_bytes(content)
        elif file_ext in ['.doc', '.docx']:
            return self._extract_from_word_bytes(content)
        elif file_ext == '.txt':
            return content.decode('utf-8', errors='ignore'), {"format": "txt"}
        else:
            return "", {"error": f"Desteklenmeyen format: {file_ext}"}
    
    def _extract_from_pdf(self, file_path: str) -> Tuple[str, dict]:
        """PDF dosyasından metin çıkar."""
        try:
            import PyPDF2
            
            text_parts = []
            metadata = {"format": "pdf", "pages": 0}
            
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                metadata["pages"] = len(reader.pages)
                
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            return "\n\n".join(text_parts), metadata
            
        except ImportError:
            return "", {"error": "PyPDF2 kurulu değil"}
        except Exception as e:
            return "", {"error": str(e)}
    
    def _extract_from_pdf_bytes(self, content: bytes) -> Tuple[str, dict]:
        """PDF bytes'tan metin çıkar."""
        try:
            import PyPDF2
            from io import BytesIO
            
            text_parts = []
            metadata = {"format": "pdf", "pages": 0}
            
            file = BytesIO(content)
            reader = PyPDF2.PdfReader(file)
            metadata["pages"] = len(reader.pages)
            
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            return "\n\n".join(text_parts), metadata
            
        except ImportError:
            return "", {"error": "PyPDF2 kurulu değil"}
        except Exception as e:
            return "", {"error": str(e)}
    
    def _extract_from_word(self, file_path: str) -> Tuple[str, dict]:
        """Word dosyasından metin çıkar."""
        try:
            import docx
            
            doc = docx.Document(file_path)
            text_parts = []
            metadata = {"format": "docx", "paragraphs": 0}
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            metadata["paragraphs"] = len(text_parts)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return "\n\n".join(text_parts), metadata
            
        except ImportError:
            return "", {"error": "python-docx kurulu değil"}
        except Exception as e:
            return "", {"error": str(e)}
    
    def _extract_from_word_bytes(self, content: bytes) -> Tuple[str, dict]:
        """Word bytes'tan metin çıkar."""
        try:
            import docx
            from io import BytesIO
            
            doc = docx.Document(BytesIO(content))
            text_parts = []
            metadata = {"format": "docx", "paragraphs": 0}
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            metadata["paragraphs"] = len(text_parts)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return "\n\n".join(text_parts), metadata
            
        except ImportError:
            return "", {"error": "python-docx kurulu değil"}
        except Exception as e:
            return "", {"error": str(e)}
    
    def _extract_from_txt(self, file_path: str) -> Tuple[str, dict]:
        """TXT dosyasından metin çıkar."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()
            return text, {"format": "txt"}
        except Exception as e:
            return "", {"error": str(e)}
    
    def _extract_from_image(self, file_path: str) -> Tuple[str, dict]:
        """
        Görüntü dosyasından metin çıkar (OCR).
        Not: Gerçek uygulamada pytesseract veya benzeri kullanılmalı.
        """
        # OCR integration would go here
        # For now, return placeholder
        return "", {
            "format": "image",
            "note": "OCR desteği henüz aktif değil",
            "suggestion": "Manuel metin girişi gerekli"
        }


class DocumentAnalyzer:
    """
    Belge analiz servisi.
    Belge türünü ve içeriğini analiz eder.
    """
    
    def __init__(self):
        self.extractor = TextExtractionService()
    
    def analyze_document(self, file_path: str) -> dict:
        """
        Belgeyi kapsamlı analiz et.
        
        Args:
            file_path: Dosya yolu
            
        Returns:
            Analiz sonuçları
        """
        # Extract text
        text, extraction_metadata = self.extractor.extract_text(file_path)
        
        result = {
            "extraction": extraction_metadata,
            "text_length": len(text),
            "word_count": len(text.split()) if text else 0,
            "is_empty": len(text.strip()) == 0,
        }
        
        if text:
            # Detect document type
            result["detected_type"] = self._detect_document_type(text)
            
            # Check for key elements
            result["has_date"] = self._check_for_date(text)
            result["has_signature_mention"] = self._check_for_signature_mention(text)
            result["has_amount"] = self._check_for_amount(text)
        
        return result
    
    def _detect_document_type(self, text: str) -> str:
        """Belge türünü tespit et."""
        text_lower = text.lower()
        
        if "tc kimlik" in text_lower or "nüfus cüzdanı" in text_lower:
            return "identity"
        elif "vergi levhası" in text_lower or "vergi dairesi" in text_lower:
            return "tax_certificate"
        elif "imza sirküleri" in text_lower or "noter" in text_lower:
            return "signature_circular"
        elif "ticaret sicil" in text_lower:
            return "trade_registry"
        elif "bilanço" in text_lower or "mali tablo" in text_lower:
            return "financial_statement"
        elif "kvkk" in text_lower or "kişisel veri" in text_lower:
            return "kvkk_consent"
        elif "sözleşme" in text_lower or "protokol" in text_lower:
            return "contract"
        else:
            return "unknown"
    
    def _check_for_date(self, text: str) -> bool:
        """Belgede tarih olup olmadığını kontrol et."""
        import re
        date_patterns = [
            r'\d{2}[./]\d{2}[./]\d{4}',  # DD/MM/YYYY or DD.MM.YYYY
            r'\d{4}[./]\d{2}[./]\d{2}',  # YYYY/MM/DD or YYYY.MM.DD
            r'\d{1,2}\s+(Ocak|Şubat|Mart|Nisan|Mayıs|Haziran|Temmuz|Ağustos|Eylül|Ekim|Kasım|Aralık)\s+\d{4}',
        ]
        
        for pattern in date_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        return False
    
    def _check_for_signature_mention(self, text: str) -> bool:
        """Belgede imza referansı olup olmadığını kontrol et."""
        signature_keywords = ['imza', 'paraf', 'onay', 'mühür', 'kaşe']
        text_lower = text.lower()
        return any(keyword in text_lower for keyword in signature_keywords)
    
    def _check_for_amount(self, text: str) -> bool:
        """Belgede tutar bilgisi olup olmadığını kontrol et."""
        import re
        amount_patterns = [
            r'\d{1,3}(?:[.,]\d{3})*(?:[.,]\d{2})?\s*(?:TL|₺|TRY)',
            r'(?:TL|₺|TRY)\s*\d{1,3}(?:[.,]\d{3})*(?:[.,]\d{2})?',
        ]
        
        for pattern in amount_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        return False



